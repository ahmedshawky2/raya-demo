# -*- coding: utf-8 -*-

from odoo import models, fields, api
from odoo.exceptions import ValidationError, UserError
import logging
from datetime import datetime
from datetime import date
import requests
from lxml import html
#from bs4 import BeautifulSoup
from bs4 import BeautifulSoup as bs
import base64
from io import StringIO
from io import BytesIO
from tika import parser
import docx
import PyPDF2

__all__ = ["doc2text"]

_logger = logging.getLogger(__name__)


class IndexAttachment(models.Model):
    _inherit = 'ir.attachment'

    is_indexed = fields.Boolean(string="Is Indexed", index=True, default=False, store=True)
    att_content = fields.Text(string="Attachment Content", store=True)
    indexed_date = fields.Datetime(string="Indexed DateTime", index=True, store=True)



    def IndexingAttachmentJob(self):

        all = self.env['ir.attachment'].search([('is_indexed', '=', False)])
        for rec in all:
            if rec[0]['mimetype'] == "application/pdf":

                soup = BytesIO()
                soup.write(base64.decodestring(rec[0]['datas']))

                #parsedPDF = parser.from_buffer(soup.getvalue())
                #text = parsedPDF["content"].encode('ascii', errors='ignore')

                pdfReader = PyPDF2.PdfFileReader(BytesIO(soup.getvalue()))
                count = pdfReader.numPages

                output = []
                for i in range(count):
                    page = pdfReader.getPage(i)
                    output.append(page.extractText())

                text = '\n'.join(output)

                rec[0]['att_content'] =  text
                rec[0]['is_indexed'] = True
                rec[0]['indexed_date'] = datetime.now()

                self.env.cr.commit()

            elif rec[0]['mimetype'] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":

                soup = BytesIO()
                soup.write(base64.decodestring(rec[0]['datas']))

                doc = docx.Document(BytesIO(soup.getvalue()))

                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                text = '\n'.join(fullText)

                rec[0]['att_content'] = text
                rec[0]['is_indexed'] = True
                rec[0]['indexed_date'] = datetime.now()

                self.env.cr.commit()


    def IndexingAttachmentSelf(self):

        soup = BytesIO()
        soup.write(base64.decodestring(self.datas))

        if self.mimetype == "application/pdf":
            _logger.info("soup !MAGED : " + str(soup))

            #parsedPDF = parser.from_buffer(soup.getvalue())
            #text = parsedPDF["content"].encode('ascii', errors='ignore')

            pdfReader = PyPDF2.PdfFileReader(BytesIO(soup.getvalue()))
            count = pdfReader.numPages

            output = []
            for i in range(count):
                page = pdfReader.getPage(i)
                output.append(page.extractText())

            text = '\n'.join(output)

            self.att_content =  text
            self.is_indexed = True
            self.indexed_date = datetime.now()

        elif self.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":

            doc = docx.Document(BytesIO(soup.getvalue()))

            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            text = '\n'.join(fullText)

            self.att_content = text
            self.is_indexed = True
            self.indexed_date = datetime.now()

            #self.env.cr.commit()
